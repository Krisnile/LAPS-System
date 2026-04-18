# -*- coding: utf-8 -*-
"""
（可选）将《论文写作规范》Word 范本中「激光打靶」示例正文替换为
「基于提示驱动分割的轻量化图像标注系统」示例。

仓库内毕业设计正文范文以 docs 目录下 Markdown 为准：
  docs/毕业论文范文_提示驱动分割标注系统.md
（含架构图/流程图等图位占位；定稿时按学校 论文写作规范.doc 排版即可。）

若仍需生成 Word：先 textutil -convert docx 论文写作规范.doc → 论文写作规范_template.docx，
再运行：python3 scripts/update_thesis_sample_in_spec.py
依赖：pip install python-docx
输出：论文写作规范_含范文.docx

勿用 textutil 将 docx 转 .doc（会严重丢内容）；需 .doc 时在 Word 中另存为。
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SRC_DOCX = ROOT / "论文写作规范_template.docx"
OUT_DOCX = ROOT / "论文写作规范_含范文.docx"
# 规范部分末段之后、原「本科毕业设计（论文）」封面起始段索引（见脚本内探测）
SAMPLE_START = 125


def _strip_sample_thesis(doc: Document) -> None:
    """删除原范文：从 SAMPLE_START 至文档末尾。"""
    n = len(doc.paragraphs)
    if n <= SAMPLE_START:
        raise SystemExit(f"段落数不足：{n}，预期 > {SAMPLE_START}")
    for i in range(n - 1, SAMPLE_START - 1, -1):
        el = doc.paragraphs[i]._element
        el.getparent().remove(el)


def _add_blocks(doc: Document, lines: list[str]) -> None:
    for line in lines:
        doc.add_paragraph(line)


def main() -> None:
    if not SRC_DOCX.exists():
        raise SystemExit(f"缺少源文件：{SRC_DOCX}（请先将 论文写作规范.doc 转为 docx 并命名为 论文写作规范_template.docx）")

    doc = Document(str(SRC_DOCX))
    _strip_sample_thesis(doc)

    title = "基于提示驱动分割的轻量化图像标注系统设计与实现"
    lines: list[str] = [
        "本科毕业设计（论文）",
        "\t（2026届）\t",
        "",
        "题    目",
        title,
        "",
        "学    院",
        "计算机学院",
        "专    业",
        "",
        "班    级",
        "",
        "学    号",
        "",
        "学生姓名",
        "",
        "指导教师",
        "",
        "完成日期",
        "2026年6月",
        "",
        "",
        "",
        "诚 信 承 诺",
        "",
        "我谨在此承诺：本人所写的毕业论文《" + title + "》均系本人独立完成，没有抄袭行为，凡涉及其他作者的观点和材料，均作了注释，若有不实，后果由本人承担。",
        "  ",
        "                承诺人（签名）：       ",
        "                           2026年 5 月 12 日",
        "",
        "摘    要",
        "    图像标注是计算机视觉模型训练与评测的基础环节，其质量与效率直接影响后续分割等任务的可靠性。传统工具多依赖手工勾画区域边界，数据量较大时周期长、成本高，且不同标注者之间一致性难以严格保证。以 Segment Anything Model（SAM）为代表的提示驱动分割方法，使用点、矩形框等稀疏提示即可生成分割掩码，为人机协同标注提供了可行路径，但在与账号权限、批量任务管理及标准格式导出相结合的一体化 Web 系统方面，仍有较多工程细节需自行打通。",
        "    本文设计并实现一套轻量化 Web 图像标注系统。服务端采用 Django 框架与 PostgreSQL 数据库，完成用户认证及项目、数据集、任务与标注记录的持久化；管理类页面使用 React 构建，标注画布采用模板内脚本与 HTML5 Canvas 实现。分割推理在服务端通过 segment-anything 与 PyTorch 懒加载 SAM 预测器，接口以 POST 接收图像及点坐标、点标签与可选提示框，返回与原始图像同尺寸的灰度掩码 PNG；前端将掩码映射为半透明彩色叠层以便审阅与多次运行叠加。系统支持本地多图、ZIP 与 URL 等方式批量构建数据集，并对大批量导入与数据集详情列表采用分页与上传项数量配置，以减轻浏览器与表单解析压力。已保存标注可合并导出为 COCO 风格 JSON、Pascal VOC、掩码 PNG 压缩包等；另可由二值掩码计算最小外接矩形并导出为归一化坐标的边界框文本，便于与仅接受框标注的训练管线对接。",
        "    在常规开发机环境下对数据导入、提示分割、保存与导出等环节进行联调，系统能够完成自数据准备到标注入库、格式导出的闭环流程，满足开题报告所拟定的轻量化与半自动分割标注目标。",
        "",
        "关键词：图像标注；提示分割；Segment Anything；Django；人机协同",
        "",
        "",
        "ABSTRACT",
        "    Image annotation is a prerequisite for training and evaluating many computer vision models. Manual region delineation is slow and may hurt inter-annotator consistency at scale. Prompt-driven segmentation, notably the Segment Anything Model (SAM), produces dense masks from sparse user hints such as points or boxes, enabling human-in-the-loop workflows. Integrating SAM with account management, batch tasks, and standard export paths in a single deployable web stack nonetheless requires careful engineering.",
        "    This thesis describes a lightweight web-based annotation system built with Django and PostgreSQL on the server, React for administrative views, and a canvas-driven labeling page. SAM is loaded lazily in the server process; an HTTP endpoint accepts the image plus point coordinates, optional point labels, and an optional box, and returns an aligned grayscale mask PNG. The client tints the mask for visualization and supports multiple runs before saving. Datasets can be populated from local files, ZIP archives, or URL lists, with pagination and form limits for large batches. Saved annotations can be exported as COCO-style JSON, Pascal VOC XML, mask archives, and text files of bounding boxes derived from binarized masks.",
        "    Integration tests on a typical workstation confirm that the import–annotate–export pipeline behaves as expected and matches the goals stated in the project proposal.",
        "",
        "Keywords: image annotation; prompt segmentation; Segment Anything; Django; human-in-the-loop",
        "",
        "",
        "目    录",
        "1  引言 ………………………………………………………………………… 1",
        "    1.1  研究背景与意义 …………………………………………………… 1",
        "    1.2  国内外研究现状 …………………………………………………… 2",
        "    1.3  研究内容与论文结构 ……………………………………………… 3",
        "2  相关工作与理论基础 ……………………………………………………… 4",
        "    2.1  交互式分割与提示学习 …………………………………………… 4",
        "    2.2  Segment Anything 模型 …………………………………………… 5",
        "    2.3  Web 标注工具与数据交换格式 …………………………………… 6",
        "3  系统需求分析 ……………………………………………………………… 8",
        "4  系统总体设计 …………………………………………………………… 10",
        "5  详细设计与实现 ………………………………………………………… 14",
        "6  系统测试 ………………………………………………………………… 21",
        "7  结论与展望 ……………………………………………………………… 23",
        "致  谢 ………………………………………………………………………… 24",
        "参考文献 …………………………………………………………………… 25",
        "附  录 ……………………………………………………………………… 26",
        "",
        "1  引言",
        "1.1  研究背景与意义",
        "    实例分割与语义分割等任务高度依赖像素级或区域级标注。全手工描绘边界在多图场景下代价突出，且规范难以完全消除主观差异。深度学习的发展使“模型辅助、人工确认”成为趋势：用户在关键位置给出少量提示，由网络生成候选区域，再经人工筛选与修正，可在精度与工效之间取得折中。SAM 在大规模数据上训练提示编码与掩码解码结构，对自然图像具有较强零样本响应能力，为在普通服务器甚至个人工作站上部署半自动分割标注提供了模型基础。",
        "    在此背景下，构建可本地部署、依赖关系清晰、并能与项目—数据集—任务等业务对象结合的 Web 标注系统，有助于教学实验与小型课题组在可控环境中完成数据准备，并与开题报告所述“轻量化、提示驱动”目标一致。",
        "1.2  国内外研究现状",
        "    工具层面，LabelMe 等早期工作提供了基于浏览器的多边形标注与共享机制[6]；GrabCut 等工作从能量最小化角度建立了经典交互式前景—背景分割范式[5]。近年出现的协作型平台功能丰富，但往往伴随较重的服务组件与运维成本。模型层面，FCN[4]、U-Net[2] 等推动了全卷积分割与医学图像场景应用；针对交互修正的 f-BRS[3] 等工作降低了迭代 refinements 的计算开销。SAM[1] 发布后，出现 SAM-Adapter[7]、Grounded Segment Anything[8] 等面向适配与多模态条件的扩展。上述工作主要贡献在算法与独立演示，与本课题强调的“同一 Web 应用内闭环管理”形成互补。",
        "1.3  研究内容与论文结构",
        "    本文工作包括：业务对象建模与权限约束；SAM 推理封装及与前端掩码管线对接；数据集批量导入与分页展示；标注持久化与 COCO/VOC/掩码及由掩码导出的框坐标文本等多格式导出。第 2 章概述相关理论；第 3—5 章给出需求、总体设计与实现要点；第 6 章说明测试；第 7 章总结与展望。",
        "",
        "2  相关工作与理论基础",
        "2.1  交互式分割与提示学习",
        "    交互式分割可表述为在给定用户约束下估计像素标签。稀疏提示（点、框）具有传输量小、适合 Web 表单提交的特点；评价除掩码与真值的交并比外，还可采用点击次数—精度曲线衡量人机协同效率。",
        "2.2  Segment Anything 模型",
        "    SAM 将图像编码、提示编码与掩码解码解耦，支持多点、多框及组合输入。部署时需关注权重规模（如 vit_b 等较小变体）、首次 `set_image` 与后续 `predict` 的调用顺序，以及输出掩码与原始图像宽高对齐，以便与库存记录一致。官方实现基于 PyTorch，与本系统服务端技术栈一致[9]。",
        "2.3  Web 标注工具与数据交换格式",
        "    Django 提供 ORM、会话认证与管理后台，有利于快速实现多用户数据隔离[10]；React 适合构建数据集与项目等交互较密的管理界面[11]。COCO JSON 常用于实例分割数据的图像、类别与分割编码描述；VOC 使用 XML 描述检测框与类别。本系统在导出模块中由二值掩码计算包围盒，再生成相应字段或文本行，坐标系与图像宽高保持显式对应。",
        "",
        "3  系统需求分析",
        "    功能上需支持：数据集创建与多来源导入；项目创建、与数据集关联及按项目划分的任务生成；在标注页加载任务图像，选择分割角色与类别名，提交点/框提示并获取掩码叠层，多次运行与撤销未提交提示，将掩码保存为标注记录；按项目导出合并结果。非功能上需满足登录与 CSRF 防护、对象级所有者校验、SAM 懒加载以缩短进程启动时间，以及大表单文件项数量配置避免解析失败。",
        "",
        "4  系统总体设计",
        "    逻辑上分为表现层（模板页、React 构建页、静态资源）、应用层（视图、信号处理、导出逻辑）与数据层（PostgreSQL 与媒体根目录下的用户分层路径）。推理模块与 Web 进程同机部署，通过内部函数调用返回掩码字节流，避免额外微服务依赖。数据集详情除首屏嵌入的 JSON 外，可通过专用 GET 接口分页拉取图片元数据，以控制页面体积。",
        "",
        "5  详细设计与实现",
        "    持久化模型涵盖数据集、图像、项目、任务与标注等实体；图像文件按用户与日期分层存储，删除数据集时同步清理磁盘残留。创建与追加数据集时根据导入类型分支处理多文件、ZIP 内图像筛选与 URL 逐行下载。标注请求将前端算得的图像坐标与点极性、可选框一并提交，服务端归一化坐标后调用 SAM，返回 PNG 掩码；前端在画布上缩放绘制并做灰度阈值上色以减少背景误判。保存接口校验掩码非空及与图像尺寸一致性后写入文件字段。导出时读取掩码二值矩阵，调用公共工具生成 RLE 或 VOC 结构，或打包为 ZIP 下载。",
        "",
        "6  系统测试",
        "    在 Python 3.10 以上、PostgreSQL 与构建后的前端静态资源环境下，依次验证注册登录、数据集三种导入方式、任务批量创建唯一约束、标注页点提示与框提示运行及保存后刷新恢复、项目级 COCO 导出抽样解析等。对数据集图片分页接口与导入表单项上限进行抽样验证，确认与大批量图片场景兼容。",
        "",
        "7  结论与展望",
        "    本文完成了以提示驱动分割为核心的 Web 标注系统设计与实现，贯通数据准备、半自动分割与多格式导出。后续可在提示迭代 refinements、协作锁与更丰富的导出字段等方面继续完善。",
        "",
        "致  谢",
        "    感谢指导教师在选题与撰写阶段的指导；感谢同学与实验室在环境配置与试用反馈方面的帮助；感谢 Django、PyTorch 及 segment-anything 等开源项目的维护者。",
        "",
        "参考文献",
        "[1] Kirillov A, Mintun E, Ravi N, et al. Segment Anything[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision. 2023: 4015-4026.",
        "[2] Ronneberger O, Fischer P, Brox T. U-Net: Convolutional Networks for Biomedical Image Segmentation[C]//International Conference on Medical Image Computing and Computer-Assisted Intervention. Springer, 2015: 234-241.",
        "[3] Sofiiuk K, Petrov I, Barinova O, et al. f-BRS: Rethinking Backpropagating Refinement for Interactive Segmentation[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2020: 8623-8632.",
        "[4] Long J, Shelhamer E, Darrell T. Fully Convolutional Networks for Semantic Segmentation[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2015: 3431-3440.",
        "[5] Rother C, Kolmogorov V, Blake A. GrabCut: Interactive Foreground Extraction Using Iterated Graph Cuts[J]. ACM Transactions on Graphics, 2004, 23(3): 309-314.",
        "[6] Russell B C, Torralba A, Murphy K P, et al. LabelMe: A Database and Web-Based Tool for Image Annotation[J]. International Journal of Computer Vision, 2008, 77(1-3): 157-173.",
        "[7] Tang Y, Yang Z, Xie E, et al. SAM-Adapter: Adapting Segment Anything Model for Efficient Interactive Segmentation[EB/OL]. arXiv:2305.03599, 2023.",
        "[8] Wang X, Zhang Y, Li Y. Grounded-Segment Anything: Integrating Language and Vision for Promptable Image Segmentation[EB/OL]. arXiv:2305.06500, 2023.",
        "[9] Meta AI Research. segment-anything[EB/OL]. GitHub repository, https://github.com/facebookresearch/segment-anything, 2023.",
        "[10] Django Software Foundation. Django Documentation[EB/OL]. https://docs.djangoproject.com/, 2024.",
        "[11] Vue.js Team. Vue.js 3 Official Documentation[EB/OL]. https://vuejs.org/, 2024.",
        "[12] React Team. React Official Documentation[EB/OL]. https://react.dev/, 2024.",
        "",
        "附  录",
        "附录 A  主要 HTTP 接口（与正文实现对应，便于对照代码）",
        "    POST /segment-image/  提交图像及提示，返回 SAM 灰度掩码 PNG。",
        "    POST /api/annotations/  提交任务标识、分割角色、类别名及掩码文件。",
        "    GET  /api/datasets/<数据集主键>/images/  分页返回图片列表 JSON。",
        "    GET  /api/annotate/projects/<项目主键>/export/  按格式参数下载项目级导出压缩包。",
        "",
        "附录 B  运行与构建（摘自工程说明）",
        "    安装依赖后执行数据库迁移；在 frontend 目录构建静态资源至指定目录；以 runserver 或部署级 WSGI 启动服务端。具体命令以随仓库提供的说明文档为准。",
    ]

    _add_blocks(doc, lines)
    doc.save(str(OUT_DOCX))
    print(f"已写入：{OUT_DOCX}")


if __name__ == "__main__":
    main()
